"""
newsletter_generator.py — Produce the final formatted newsletter.
Outputs:
  • Plain-text  (.txt)
  • Markdown    (.md)
  • HTML        (.html)
  • Word Doc    (.docx)
"""

import os
from datetime import datetime
from collections import defaultdict

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from config import CATEGORY_ORDER
from utils.logger import get_logger
from utils.models import Article

logger = get_logger(__name__)


def _group_by_category(articles: list[Article]) -> dict:
    groups = defaultdict(list)
    for art in articles:
        category = art.category if art.category else "General Finance"
        groups[category].append(art)
    return groups


# ── Plain-text renderer ───────────────────────────────────────────────────────
def render_text(articles: list[Article], week_ending: str) -> str:
    lines = [
        "=" * 65,
        "  HFC WEEKLY FINANCIAL INTELLIGENCE NEWSLETTER",
        f"  Week Ending: {week_ending}",
        "=" * 65,
        "",
    ]

    groups = _group_by_category(articles)
    for category in CATEGORY_ORDER:
        arts = groups.get(category, [])
        if not arts:
            continue
        lines += [
            f"{'─' * 65}",
            f"  {category.upper()}",
            f"{'─' * 65}",
            "",
        ]
        for art in arts:
            pub = art.pub_date.strftime("%d %b %Y") if art.pub_date else "Date unknown"
            lines += [
                f"• {art.title}",
                f"  Source: {art.source}  |  {pub}",
                f"  {art.summary}" if art.summary else "",
                f"  ➤ Key Takeaway: {art.takeaway}" if art.takeaway else "",
                f"  🔗 {art.url}",
                "",
            ]

    lines += [
        "=" * 65,
        "  HFC Group Plc — Internal Distribution Only",
        f"  Generated: {datetime.now().strftime('%d %B %Y %H:%M')}",
        "=" * 65,
    ]
    return "\n".join(lines)


# ── Markdown renderer ─────────────────────────────────────────────────────────
def render_markdown(articles: list[Article], week_ending: str) -> str:
    lines = [
        "# HFC Weekly Financial Intelligence Newsletter",
        f"**Week Ending:** {week_ending}",
        "---",
        ""
    ]

    groups = _group_by_category(articles)
    for category in CATEGORY_ORDER:
        arts = groups.get(category, [])
        if not arts:
            continue
        lines.append(f"## {category}")
        for art in arts:
            pub = art.pub_date.strftime("%d %b %Y") if art.pub_date else "Date unknown"
            lines.append(f"### [{art.title}]({art.url})")
            lines.append(f"**Source:** {art.source} | **Date:** {pub}")
            if art.summary:
                lines.append(f"\n{art.summary}")
            if art.takeaway:
                lines.append(f"\n> **Key Takeaway:** {art.takeaway}")
            lines.append("\n---")

    lines.append("\n*HFC Group Plc — Internal Distribution Only*")
    lines.append(f"*Generated: {datetime.now().strftime('%d %B %Y %H:%M')}*")
    return "\n".join(lines)


# ── HTML renderer ─────────────────────────────────────────────────────────────
def render_html(articles: list[Article], week_ending: str) -> str:
    groups = _group_by_category(articles)

    section_html = ""
    for category in CATEGORY_ORDER:
        arts = groups.get(category, [])
        if not arts:
            continue
        items_html = ""
        for art in arts:
            pub = art.pub_date.strftime("%d %b %Y") if art.pub_date else "Date unknown"
            summary_html = f'<p class="summary">{art.summary}</p>' if art.summary else ""
            takeaway_html = f'<p class="takeaway"><strong>Key Takeaway:</strong> {art.takeaway}</p>' if art.takeaway else ""
            items_html += f"""
            <div class="article">
              <h3><a href="{art.url}" target="_blank">{art.title}</a></h3>
              <p class="meta">{art.source} &nbsp;|&nbsp; {pub}</p>
              {summary_html}
              {takeaway_html}
            </div>"""
        section_html += f"""
        <section>
          <h2>{category}</h2>
          {items_html}
        </section>"""

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>HFC Weekly Newsletter — {week_ending}</title>
<style>
  body {{ font-family: Georgia, serif; max-width: 800px; margin: 0 auto; padding: 20px; color: #222; margin-bottom: 50px; }}
  header {{ background: #003366; color: white; padding: 24px; border-radius: 4px; margin-bottom: 24px; }}
  header h1 {{ margin: 0; font-size: 1.6em; }}
  header p {{ margin: 4px 0 0; opacity: .8; }}
  section {{ margin-bottom: 32px; }}
  h2 {{ background: #e8f0fe; padding: 8px 14px; border-left: 4px solid #003366; font-size: 1.1em; }}
  .article {{ border-bottom: 1px solid #eee; padding: 12px 0; }}
  .article h3 {{ margin: 0 0 4px; font-size: 1em; }}
  .article h3 a {{ color: #003366; text-decoration: none; }}
  .meta {{ font-size: .8em; color: #777; margin: 0 0 6px; }}
  .summary {{ margin: 0 0 4px; }}
  .takeaway {{ font-size: .9em; background: #fffbe6; padding: 6px 10px; border-radius: 3px; border-left: 3px solid #ffcc00; }}
  footer {{ font-size: .8em; color: #888; border-top: 1px solid #ddd; padding-top: 12px; margin-top: 32px; text-align: center; }}
</style>
</head>
<body>
<header>
  <h1>HFC Weekly Financial Intelligence Newsletter</h1>
  <p>Week Ending: {week_ending}</p>
</header>
{section_html}
<footer>
  <p>HFC Group Plc — Internal Distribution Only</p>
  <p>Generated: {datetime.now().strftime("%d %B %Y %H:%M")}</p>
</footer>
</body>
</html>"""


# ── Word Docx renderer ────────────────────────────────────────────────────────
def export_docx(articles: list[Article], week_ending: str, docx_path: str) -> None:
    doc = Document()
    
    # Title
    title = doc.add_heading('HFC Weekly Financial Intelligence Newsletter', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(f"Week Ending: {week_ending}")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Content
    groups = _group_by_category(articles)
    for category in CATEGORY_ORDER:
        arts = groups.get(category, [])
        if not arts:
            continue
            
        cat_heading = doc.add_heading(category, level=1)
        
        for art in arts:
            pub = art.pub_date.strftime("%d %b %Y") if art.pub_date else "Date unknown"
            
            # Title
            p_title = doc.add_heading(art.title, level=2)
            
            # Metadata
            p_meta = doc.add_paragraph()
            p_meta.add_run(f"Source: {art.source} | Date: {pub}").italic = True
            
            # Summary
            if art.summary:
                p_summary = doc.add_paragraph(art.summary)
            
            # Takeaway
            if art.takeaway:
                p_takeaway = doc.add_paragraph()
                p_takeaway.add_run("Key Takeaway: ").bold = True
                p_takeaway.add_run(art.takeaway)
                
            # URL
            p_url = doc.add_paragraph()
            p_url.add_run("Read more: ")
            p_url.add_run(art.url).underline = True
            
            doc.add_paragraph() # spacing

    doc.add_paragraph("=" * 50)
    p_footer = doc.add_paragraph("HFC Group Plc — Internal Distribution Only\n")
    p_footer.add_run(f"Generated: {datetime.now().strftime('%d %B %Y %H:%M')}")
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(docx_path)
    logger.info("Word document saved to %s", docx_path)


# ── Generator ─────────────────────────────────────────────────────────────────
def generate_newsletter(articles: list[Article], output_dir: str = "output") -> dict[str, str]:
    """Run all renderers, write files, return dict of output paths."""
    os.makedirs(output_dir, exist_ok=True)
    date_str = datetime.now().strftime("%Y-%m-%d")
    week_ending = datetime.now().strftime("%d %B %Y")

    txt_path  = os.path.join(output_dir, f"hfc_newsletter_{date_str}.txt")
    md_path   = os.path.join(output_dir, f"hfc_newsletter_{date_str}.md")
    html_path = os.path.join(output_dir, f"hfc_newsletter_{date_str}.html")
    docx_path = os.path.join(output_dir, f"hfc_newsletter_{date_str}.docx")

    txt_content  = render_text(articles, week_ending)
    md_content   = render_markdown(articles, week_ending)
    html_content = render_html(articles, week_ending)

    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(txt_content)
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html_content)
        
    export_docx(articles, week_ending, docx_path)

    logger.info("Newsletter saved to: \n • %s\n • %s\n • %s\n • %s", txt_path, md_path, html_path, docx_path)
    return {
        "txt": txt_path,
        "md": md_path,
        "html": html_path,
        "docx": docx_path
    }
